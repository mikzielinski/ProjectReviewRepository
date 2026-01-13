#!/usr/bin/env python3
"""Analyze DOCX file structure"""
import sys
import os

# Add backend to path
sys.path.insert(0, os.path.join(os.path.dirname(__file__), 'backend'))

try:
    from docx import Document
    from docx.oxml.ns import qn
    
    docx_path = '/Users/mikolajzielinski/Downloads/PDD Template_PDD.docx'
    
    if not os.path.exists(docx_path):
        print(f"File not found: {docx_path}")
        sys.exit(1)
    
    doc = Document(docx_path)
    
    print("=" * 80)
    print("DOCX FILE ANALYSIS")
    print("=" * 80)
    print(f"Total paragraphs: {len(doc.paragraphs)}")
    print(f"Total tables: {len(doc.tables)}")
    print(f"Total sections: {len(doc.sections)}")
    
    print("\n" + "=" * 80)
    print("SECTION ORIENTATIONS")
    print("=" * 80)
    for i, section in enumerate(doc.sections):
        width = section.page_width
        height = section.page_height
        is_landscape = width > height
        orientation = "landscape" if is_landscape else "portrait"
        print(f"Section {i}: {width} x {height} EMU -> {orientation}")
        if "Software" in str(section) or i == 1:  # Check section 1 (0-indexed)
            print(f"  ^ This section should be landscape for 'Software Preparation'")
    
    print("\n" + "=" * 80)
    print("FIRST 30 PARAGRAPHS")
    print("=" * 80)
    for i, p in enumerate(doc.paragraphs[:30]):
        text = p.text[:100].replace('\n', ' ') if p.text else '(empty)'
        style = p.style.name if p.style else 'None'
        has_runs = len(p.runs) > 0
        has_images = False
        if has_runs:
            for run in p.runs:
                if hasattr(run, '_element'):
                    inline_pics = run._element.findall('.//{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}inline')
                    if inline_pics:
                        has_images = True
                        break
        status = ""
        if not text or text == '(empty)':
            if has_images:
                status = " [HAS IMAGE]"
            else:
                status = " [EMPTY]"
        print(f"{i:3d}: [{style:20s}] {text}{status}")
    
    print("\n" + "=" * 80)
    print("TABLES ANALYSIS")
    print("=" * 80)
    for i, table in enumerate(doc.tables):
        print(f"\nTable {i}: {len(table.rows)} rows x {len(table.columns)} cols")
        # Check for images in cells
        images_found = []
        for row_idx, row in enumerate(table.rows[:10]):  # Check first 10 rows
            for cell_idx, cell in enumerate(row.cells[:5]):  # Check first 5 cols
                has_image = False
                for para in cell.paragraphs:
                    for run in para.runs:
                        if hasattr(run, '_element'):
                            inline_pics = run._element.findall('.//{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}inline')
                            if inline_pics:
                                has_image = True
                                images_found.append((row_idx, cell_idx))
                                break
                    if has_image:
                        break
                if has_image:
                    break
        if images_found:
            print(f"  Images found in cells: {images_found}")
        else:
            print(f"  No images found in first 10 rows x 5 cols")
    
    print("\n" + "=" * 80)
    print("CHECKING FOR 'Software Preparation' AND 'As Is Process Map'")
    print("=" * 80)
    software_found = False
    process_map_found = False
    for i, p in enumerate(doc.paragraphs):
        text = p.text if p.text else ""
        if "Software Preparation" in text:
            print(f"Found 'Software Preparation' at paragraph {i}: {text[:100]}")
            software_found = True
        if "As Is Process Map" in text or "Program Logic" in text:
            print(f"Found 'As Is Process Map/Program Logic' at paragraph {i}: {text[:100]}")
            process_map_found = True
    
    if not software_found:
        print("WARNING: 'Software Preparation' not found in paragraphs!")
    if not process_map_found:
        print("WARNING: 'As Is Process Map' or 'Program Logic' not found in paragraphs!")
    
    print("\n" + "=" * 80)
    print("DOCUMENT BODY ELEMENTS (XML)")
    print("=" * 80)
    body = doc.element.body
    elements = list(body)
    print(f"Total XML elements in body: {len(elements)}")
    
    # Count element types
    element_types = {}
    for elem in elements:
        tag = elem.tag.split('}')[-1] if '}' in elem.tag else elem.tag
        element_types[tag] = element_types.get(tag, 0) + 1
    
    print("Element type counts:")
    for tag, count in sorted(element_types.items()):
        print(f"  {tag}: {count}")
    
    # Check for section breaks
    sect_breaks = 0
    for elem in elements:
        if elem.tag.endswith('}sectPr'):
            sect_breaks += 1
        elif elem.tag.endswith('}p'):
            sect_pr = elem.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}sectPr')
            if sect_pr is not None:
                sect_breaks += 1
    
    print(f"\nSection breaks found: {sect_breaks}")
    
except ImportError as e:
    print(f"Import error: {e}")
    print("Make sure python-docx is installed: pip install python-docx")
except Exception as e:
    print(f"Error: {e}")
    import traceback
    traceback.print_exc()


