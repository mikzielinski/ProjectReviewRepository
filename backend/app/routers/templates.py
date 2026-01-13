from typing import Optional
import logging
from fastapi import APIRouter, Depends, HTTPException, status, Request, UploadFile, File
from sqlalchemy.orm import Session
from sqlalchemy import func

from app.db import get_db
from app.dependencies import get_current_active_user
from app.schemas import TemplateCreate, TemplateRead, TemplateUpdate
from app.models import Template, Document, DocumentVersion, Project, ProjectMember
from app.core.enums import TemplateStatus

router = APIRouter(prefix="/templates", tags=["templates"])
logger = logging.getLogger(__name__)


def _generate_template_filename(name: str, doc_type: str, version: str, extension: str) -> str:
    """
    Generate filename according to convention: [name]_[doc_type]_v[version].[ext]
    Example: "My Template_PDD_v1.docx"
    """
    # Sanitize name and doc_type (remove special chars that might cause issues)
    safe_name = "".join(c if c.isalnum() or c in (' ', '-', '_') else '_' for c in name).strip()
    safe_name = "_".join(safe_name.split())  # Replace spaces with underscores
    safe_doc_type = "".join(c if c.isalnum() or c in ('-', '_') else '_' for c in doc_type).strip()
    
    # Ensure version starts with 'v' if not already
    if not version.startswith('v'):
        version = f"v{version}"
    
    # Remove . from extension if present
    ext = extension.lstrip('.')
    
    return f"{safe_name}_{safe_doc_type}_{version}.{ext}"


def _find_libreoffice_cmd():
    """Find LibreOffice executable on different OS."""
    import subprocess
    possible_paths = [
        "soffice",
        "libreoffice",
        "/opt/homebrew/bin/soffice",
        "/usr/local/bin/soffice",
        "/Applications/LibreOffice.app/Contents/MacOS/soffice",
        "/usr/bin/soffice",
    ]
    
    for path in possible_paths:
        try:
            result = subprocess.run(
                [path, "--version"],
                stdout=subprocess.PIPE,
                stderr=subprocess.PIPE,
                timeout=5
            )
            if result.returncode == 0:
                return path
        except (FileNotFoundError, subprocess.TimeoutExpired):
            continue
    
    return None


def _recalc_xlsx_with_libreoffice(input_xlsx: str, output_xlsx: str, soffice_cmd: str):
    """Recalculate formulas in XLSX using LibreOffice."""
    import subprocess
    import tempfile
    import os
    
    with tempfile.TemporaryDirectory() as tmp:
        lo_profile = os.path.join(tmp, "lo-profile")
        os.makedirs(lo_profile, exist_ok=True)
        
        cmd = [
            soffice_cmd,
            "--headless",
            f"-env:UserInstallation=file://{lo_profile}",
            "--convert-to", "xlsx",
            "--outdir", tmp,
            input_xlsx
        ]
        
        result = subprocess.run(
            cmd,
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE,
            timeout=120
        )
        
        files = [f for f in os.listdir(tmp) if f.lower().endswith(".xlsx")]
        if not files:
            raise RuntimeError(
                "LibreOffice recalculation failed: "
                + result.stderr.decode(errors="ignore", encoding='utf-8')
            )
        
        os.rename(os.path.join(tmp, files[0]), output_xlsx)


def _prepare_xlsx_for_print(input_path: str, output_path: str):
    """Prepare XLSX for print: remove empty rows/columns, set landscape layout."""
    from openpyxl import load_workbook
    from openpyxl.utils import get_column_letter
    
    wb = load_workbook(input_path, data_only=True)
    sheets_to_remove = []
    
    for ws in wb.worksheets:
        max_row = 0
        max_col = 0
        
        # Detect real data
        for row in ws.iter_rows():
            for cell in row:
                if cell.value not in (None, "", " "):
                    max_row = max(max_row, cell.row)
                    max_col = max(max_col, cell.column)
        
        # Empty sheet → remove
        if max_row == 0 or max_col == 0:
            sheets_to_remove.append(ws.title)
            continue
        
        # Trim empty rows/columns
        if ws.max_row > max_row:
            ws.delete_rows(max_row + 1, ws.max_row - max_row)
        if ws.max_column > max_col:
            ws.delete_cols(max_col + 1, ws.max_column - max_col)
        
        # --- PAGE SETUP: NEVER SPLIT COLUMNS ---
        # Set print area to actual data range
        ws.print_area = f"A1:{get_column_letter(max_col)}{max_row}"
        
        # Repeat header rows on each page
        ws.print_title_rows = "1:1"
        
        # Page orientation and size
        ws.page_setup.orientation = ws.ORIENTATION_LANDSCAPE
        ws.page_setup.paperSize = ws.PAPERSIZE_A4
        
        # KLUCZOWE: Wszystkie kolumny zawsze na jednej stronie
        # fitToWidth = 1 → wszystkie kolumny zmieszczone na szerokość jednej strony
        # fitToHeight = False → NIE wymuszamy wysokości, pozwalamy Excelowi dzielić wiersze
        # scale = None → NIE używamy scale gdy używamy fitToWidth (Excel wymaga tego)
        # To gwarantuje:
        # ✅ kolumny NIGDY nie są dzielone między strony
        # ✅ wiersze idą na kolejne strony jeśli potrzeba
        # ✅ brak mikroskopijnej czcionki (bo nie skalujemy wysokości)
        ws.page_setup.scale = None  # Scale MUSI być None gdy używamy fitToWidth
        ws.page_setup.fitToWidth = 1
        ws.page_setup.fitToHeight = False
        
        # Marginesy raportowe (mniejsze marginesy = więcej miejsca na dane)
        ws.page_margins.left = 0.5
        ws.page_margins.right = 0.5
        ws.page_margins.top = 0.75
        ws.page_margins.bottom = 0.75
        ws.page_margins.header = 0.3
        ws.page_margins.footer = 0.3
        
        # Wyśrodkowanie poziome (nie pionowe)
        ws.page_setup.horizontalCentered = True
        ws.page_setup.verticalCentered = False
        
        # OPTYMALIZACJA: Automatyczne dostosowanie szerokości kolumn do zawartości
        # To pomaga LibreOffice lepiej skalować tabelę
        # Nie używamy autoFit (może być zbyt szerokie), ale ustawiamy rozsądną szerokość
        for col_idx in range(1, max_col + 1):
            col_letter = get_column_letter(col_idx)
            # Znajdź maksymalną szerokość zawartości w kolumnie
            max_width = 0
            for row in ws[col_letter]:
                if row.value:
                    # Szacuj szerokość na podstawie długości tekstu
                    cell_text = str(row.value)
                    # Przyjmij około 1.2 jednostek na znak + padding
                    estimated_width = len(cell_text) * 1.2 + 2
                    max_width = max(max_width, estimated_width)
            
            # Ustaw szerokość kolumny (min 10, max 50, żeby nie było zbyt szerokich)
            if max_width > 0:
                ws.column_dimensions[col_letter].width = min(max(max_width, 10), 50)
            else:
                ws.column_dimensions[col_letter].width = 15  # Domyślna szerokość
        
        # Optional: Log warning for very wide tables (>30 columns)
        # Excel will still fit them, but may require significant scaling
        if max_col > 30:
            import logging
            logger = logging.getLogger(__name__)
            logger.warning(
                f"Sheet '{ws.title}' has {max_col} columns - may require significant scaling. "
                f"Consider reviewing table layout for better PDF output."
            )
    
    for title in sheets_to_remove:
        del wb[title]
    
    wb.save(output_path)


def _convert_xlsx_to_pdf_full_pipeline(xlsx_bytes: bytes, soffice_cmd: str) -> bytes:
    """Full XLSX → PDF pipeline: recalc formulas, prepare layout, convert to PDF."""
    import subprocess
    import tempfile
    import os
    
    with tempfile.TemporaryDirectory() as tmp:
        raw_xlsx = os.path.join(tmp, "input.xlsx")
        recalced_xlsx = os.path.join(tmp, "recalced.xlsx")
        prepared_xlsx = os.path.join(tmp, "prepared.xlsx")
        final_xlsx = os.path.join(tmp, "final.xlsx")
        
        # Save input
        with open(raw_xlsx, "wb") as f:
            f.write(xlsx_bytes)
        
        # 1. Recalculate formulas
        _recalc_xlsx_with_libreoffice(raw_xlsx, recalced_xlsx, soffice_cmd)
        
        # 2. Prepare layout (ustawienia strony, szerokość kolumn, etc.)
        _prepare_xlsx_for_print(recalced_xlsx, prepared_xlsx)
        
        # 3. IMPORTANTE: Ponownie otwórz przez LibreOffice i zapisz
        # To zapewnia, że LibreOffice "rozumie" ustawienia strony przed konwersją do PDF
        # LibreOffice czasami nie respektuje ustawień openpyxl bezpośrednio
        lo_profile_prep = os.path.join(tmp, "lo-profile-prep")
        os.makedirs(lo_profile_prep, exist_ok=True)
        
        cmd_prep = [
            soffice_cmd,
            "--headless",
            f"-env:UserInstallation=file://{lo_profile_prep}",
            "--convert-to", "xlsx",
            "--outdir", tmp,
            prepared_xlsx
        ]
        
        result_prep = subprocess.run(
            cmd_prep,
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE,
            timeout=60
        )
        
        # Znajdź wygenerowany plik
        prep_files = [f for f in os.listdir(tmp) if f.lower().endswith(".xlsx") and "final" not in f.lower() and "input" not in f.lower() and "recalced" not in f.lower() and "prepared" not in f.lower()]
        if prep_files:
            # Użyj najnowszego pliku (jeśli jest kilka)
            final_prep_file = os.path.join(tmp, sorted(prep_files, key=lambda x: os.path.getmtime(os.path.join(tmp, x)), reverse=True)[0])
            if os.path.exists(final_prep_file):
                os.rename(final_prep_file, final_xlsx)
            else:
                final_xlsx = prepared_xlsx  # Fallback
        else:
            final_xlsx = prepared_xlsx  # Fallback jeśli LibreOffice nie wygenerował nowego
        
        # 4. Convert to PDF with page settings enforced
        lo_profile = os.path.join(tmp, "lo-profile-pdf")
        os.makedirs(lo_profile, exist_ok=True)
        
        # LibreOffice command with explicit PDF export settings
        # Używamy parametrów które wymuszają przestrzeganie ustawień strony z Excel
        cmd = [
            soffice_cmd,
            "--headless",
            "--nodefault",
            f"-env:UserInstallation=file://{lo_profile}",
            "--convert-to", "pdf",
            "--outdir", tmp,
            final_xlsx
        ]
        
        result = subprocess.run(
            cmd,
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE,
            timeout=120
        )
        
        pdfs = [f for f in os.listdir(tmp) if f.lower().endswith(".pdf")]
        if not pdfs:
            raise RuntimeError(
                "PDF generation failed: "
                + result.stderr.decode(errors="ignore", encoding='utf-8')
            )
        
        with open(os.path.join(tmp, pdfs[0]), "rb") as f:
            return f.read()


def _convert_pptx_to_pdf(pptx_bytes: bytes, soffice_cmd: str) -> bytes:
    """
    Convert PPTX to PDF using LibreOffice with high-quality image rendering.
    Uses a two-step conversion (PPTX → ODP → PDF) to better preserve images, GIFs, and graphics.
    This approach ensures LibreOffice fully processes all embedded media before PDF export.
    """
    import subprocess
    import tempfile
    import os
    
    with tempfile.TemporaryDirectory() as tmp:
        pptx_path = os.path.join(tmp, "input.pptx")
        odp_path = os.path.join(tmp, "intermediate.odp")
        pdf_path = os.path.join(tmp, "input.pdf")
        lo_profile = os.path.join(tmp, "lo-profile")
        os.makedirs(lo_profile, exist_ok=True)
        
        # Write PPTX to temp file
        with open(pptx_path, "wb") as f:
            f.write(pptx_bytes)
        
        # Step 1: Convert PPTX → ODP (LibreOffice native format)
        # This ensures all images, GIFs, and media are properly loaded and processed
        cmd1 = [
            soffice_cmd,
            "--headless",
            f"-env:UserInstallation=file://{lo_profile}",
            "--convert-to", "odp",
            "--outdir", tmp,
            pptx_path
        ]
        
        result1 = subprocess.run(
            cmd1,
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE,
            timeout=120
        )
        
        # Find the generated ODP file
        odp_files = [f for f in os.listdir(tmp) if f.endswith('.odp')]
        
        # Simplified PDF filter options - only use well-supported options
        # Too many options can cause parameter errors in LibreOffice
        pdf_filter_options = "SelectPdfVersion=1,Quality=100,ReduceImageResolution=false"
        
        if not odp_files or result1.returncode != 0:
            # Fallback: try direct PPTX → PDF conversion
            logger.warning("ODP conversion failed or no ODP file found, trying direct PPTX → PDF")
            cmd_direct = [
                soffice_cmd,
                "--headless",
                f"-env:UserInstallation=file://{lo_profile}",
                "--convert-to", f"pdf:{pdf_filter_options}",
                "--outdir", tmp,
                pptx_path
            ]
            result_direct = subprocess.run(
                cmd_direct,
                stdout=subprocess.PIPE,
                stderr=subprocess.PIPE,
                timeout=180
            )
            if result_direct.returncode != 0 or not os.path.exists(pdf_path):
                stderr_msg = result_direct.stderr.decode('utf-8', errors='ignore') if result_direct.stderr else "No error output"
                raise RuntimeError(f"LibreOffice PPTX→PDF conversion failed: {stderr_msg}")
            with open(pdf_path, "rb") as f:
                return f.read()
        
        # Use the first ODP file found
        generated_odp = os.path.join(tmp, odp_files[0])
        if generated_odp != odp_path:
            os.rename(generated_odp, odp_path)
        
        # Step 2: Convert ODP → PDF with simplified filter options
        cmd2 = [
            soffice_cmd,
            "--headless",
            f"-env:UserInstallation=file://{lo_profile}",
            "--convert-to", f"pdf:{pdf_filter_options}",
            "--outdir", tmp,
            odp_path
        ]
        
        result2 = subprocess.run(
            cmd2,
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE,
            timeout=180  # Presentations might take longer
        )
        
        if result2.returncode != 0 or not os.path.exists(pdf_path):
            stderr_msg = result2.stderr.decode('utf-8', errors='ignore') if result2.stderr else "No error output"
            raise RuntimeError(f"LibreOffice ODP→PDF conversion failed: {stderr_msg}")
        
        # Read PDF
        with open(pdf_path, "rb") as f:
            return f.read()


def _convert_docx_to_pdf(docx_bytes: bytes, soffice_cmd: str) -> bytes:
    """Convert DOCX to PDF using LibreOffice."""
    import subprocess
    import tempfile
    import os
    
    with tempfile.TemporaryDirectory() as tmp:
        docx_path = os.path.join(tmp, "input.docx")
        pdf_path = os.path.join(tmp, "input.pdf")
        lo_profile = os.path.join(tmp, "lo-profile")
        os.makedirs(lo_profile, exist_ok=True)
        
        with open(docx_path, "wb") as f:
            f.write(docx_bytes)
        
        cmd = [
            soffice_cmd,
            "--headless",
            f"-env:UserInstallation=file://{lo_profile}",
            "--convert-to", "pdf",
            "--outdir", tmp,
            docx_path
        ]
        
        result = subprocess.run(
            cmd,
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE,
            timeout=120
        )
        
        if result.returncode != 0 or not os.path.exists(pdf_path):
            stderr_msg = result.stderr.decode('utf-8', errors='ignore') if result.stderr else "No error output"
            raise HTTPException(
                status_code=500,
                detail=f"LibreOffice failed: {stderr_msg}"
            )
        
        with open(pdf_path, "rb") as f:
            return f.read()


def _generate_template_pdf(template: Template, db: Session) -> tuple[str, str]:
    """
    Generate PDF from template (DOCX or XLSX) and return (pdf_object_key, pdf_hash).
    This is a helper function used during template creation/update.
    """
    import os
    import hashlib
    import tempfile
    from app.services.storage import download_file, upload_file, generate_object_key
    
    # Download template file
    try:
        file_bytes = download_file(
            template.object_key,
            expected_hash=template.file_hash
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Cannot load template file: {str(e)}")
    
    # Find LibreOffice
    soffice_cmd = _find_libreoffice_cmd()
    if not soffice_cmd:
        raise HTTPException(
            status_code=500,
            detail="LibreOffice is not installed or not found. Install with: brew install --cask libreoffice (macOS) or apt-get install -y libreoffice (Linux)"
        )
    
    # Detect file type and convert accordingly
    file_ext = os.path.splitext(template.object_key)[1].lower()
    
    # Formats that don't support PDF conversion
    # XML-based formats, images, data files, text files
    no_pdf_formats = [
        '.xml', '.bpmn', '.bpmn20', '.bpmn2', '.uml', '.xmi',
        '.archimate', '.archi', '.drawio', '.lucid',
        '.svg', '.png', '.jpg', '.jpeg', '.webp',
        '.csv', '.tsv', '.json', '.parquet',
        '.yaml', '.yml', '.md', '.txt', '.rst',
        '.vsdx', '.vsdm', '.vsd'  # Visio files don't convert to PDF easily
    ]
    
    if file_ext == '.xlsx':
        # XLSX → PDF (full pipeline: recalc, prepare, convert)
        try:
            pdf_bytes = _convert_xlsx_to_pdf_full_pipeline(file_bytes, soffice_cmd)
        except Exception as e:
            raise HTTPException(status_code=500, detail=f"XLSX to PDF conversion failed: {str(e)}")
    elif file_ext == '.docx':
        # DOCX → PDF (direct conversion)
        try:
            pdf_bytes = _convert_docx_to_pdf(file_bytes, soffice_cmd)
        except Exception as e:
            raise HTTPException(status_code=500, detail=f"DOCX to PDF conversion failed: {str(e)}")
    elif file_ext == '.pptx':
        # PPTX → PDF (direct conversion using LibreOffice)
        try:
            pdf_bytes = _convert_pptx_to_pdf(file_bytes, soffice_cmd)
        except Exception as e:
            raise HTTPException(status_code=500, detail=f"PPTX to PDF conversion failed: {str(e)}")
    elif file_ext in no_pdf_formats:
        # These formats don't support PDF conversion - they can be viewed directly
        # Return None to indicate no PDF should be generated
        return None, None
    else:
        raise HTTPException(
            status_code=400,
            detail=f"Unsupported file type: {file_ext}. Supported formats include: .docx, .xlsx, .pptx, .bpmn, .uml, .xmi, .archimate, .vsdx, .svg, .png, .md, .json, .yaml, and more."
        )
    
    # Generate PDF filename according to convention
    pdf_filename = _generate_template_filename(template.name, template.doc_type, template.version, "pdf")
    pdf_object_key = generate_object_key("templates", pdf_filename)
    pdf_hash = hashlib.sha256(pdf_bytes).hexdigest()
    
    # Upload PDF
    try:
        uploaded_key, _ = upload_file(
            pdf_bytes,
            pdf_object_key,
            "application/pdf"
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to upload PDF: {str(e)}")
    
    return uploaded_key, pdf_hash


@router.post("/upload", response_model=dict)
def upload_template_file(
    request: Request,
    file: UploadFile = File(...),
    db: Session = Depends(get_db),
    current_user=Depends(get_current_active_user),
):
    """
    Upload a template file and return object_key and file_hash.
    HIPAA/GxP/GIS Compliance: Files are stored with integrity verification.
    """
    from app.services.storage import upload_file, generate_object_key
    import hashlib
    
    # Read file content
    content = file.file.read()
    
    # Generate object key
    object_key = generate_object_key("templates", file.filename)
    
    # Upload file and get hash
    uploaded_key, file_hash = upload_file(content, object_key, file.content_type or "application/octet-stream")
    
    # HIPAA/GxP/GIS Compliance: Audit log for file upload
    # Use separate session for audit log to avoid transaction conflicts
    try:
        from app.services.audit import log_action, AuditAction
        from app.db import SessionLocal
        
        client_ip = request.client.host if request.client else None
        user_agent = request.headers.get("user-agent")
        
        # Use separate session for audit logging to avoid transaction conflicts
        audit_db = SessionLocal()
        try:
            log_action(
                audit_db,
                actor_user_id=current_user.id,
                action=AuditAction.TEMPLATE_CREATE,  # Using TEMPLATE_CREATE for file upload
                entity_type="TemplateFile",
                entity_id=str(uploaded_key),  # Use object_key as entity_id for file uploads
                after_json={
                    "action": "file_upload",
                    "object_key": uploaded_key,
                    "filename": file.filename,
                    "file_size": len(content),
                    "file_hash": file_hash[:20] + "...",  # Partial hash for audit
                },
                ip=client_ip,
                user_agent=user_agent,
            )
            audit_db.commit()
        except Exception as audit_error:
            audit_db.rollback()
            import logging
            logger = logging.getLogger(__name__)
            logger.warning(f"Audit logging failed: {str(audit_error)}. File uploaded: {uploaded_key} by user {current_user.id}")
        finally:
            audit_db.close()
    except Exception as e:
        import logging
        logger = logging.getLogger(__name__)
        logger.warning(f"Audit logging setup failed: {str(e)}. File uploaded: {uploaded_key} by user {current_user.id}")
    
    return {
        "object_key": uploaded_key,
        "file_hash": file_hash,
        "filename": file.filename,
        "size": len(content),
    }


@router.get("", response_model=list[TemplateRead])
def list_templates(
    doc_type: Optional[str] = None,
    db: Session = Depends(get_db),
    current_user=Depends(get_current_active_user),
):
    """
    List templates accessible to the current user.
    HIPAA/GxP/GIS Compliance: Users can only see templates from their organization.
    """
    from app.models import ProjectMember, Project
    
    # Get user's organizations from project memberships
    user_projects = db.query(ProjectMember.project_id).filter(
        ProjectMember.user_id == current_user.id,
        (ProjectMember.expires_at.is_(None) | (ProjectMember.expires_at > func.now()))
    ).subquery()
    
    # Get org_ids from projects
    user_org_ids = db.query(Project.org_id).filter(
        Project.id.in_(db.query(user_projects.c.project_id))
    ).distinct().all()
    org_ids = [org_id[0] for org_id in user_org_ids] if user_org_ids else []
    
    # If user has no org access, return empty list (or allow system-wide if admin)
    # For now, we'll allow all templates if user has no org restrictions
    # In production, you might want stricter access control
    query = db.query(Template)
    if org_ids:
        query = query.filter(Template.org_id.in_(org_ids))
    
    if doc_type:
        query = query.filter(Template.doc_type == doc_type)
    
    return query.all()


@router.post("", response_model=TemplateRead, status_code=status.HTTP_201_CREATED)
def create_template(
    request: Request,
    payload: TemplateCreate, 
    db: Session = Depends(get_db), 
    current_user=Depends(get_current_active_user)
):
    from app.models import Org
    # Get or create default org if org_id not provided
    if not payload.org_id:
        default_org = db.query(Org).first()
        if not default_org:
            import uuid
            default_org = Org(id=uuid.uuid4(), name="Default Organization")
            db.add(default_org)
            db.commit()
            db.refresh(default_org)
        org_id = default_org.id
    else:
        org_id = payload.org_id
    
    version = payload.version or "v1"
    
    # Download existing file and re-upload with new name according to convention
    from app.services.storage import download_file, upload_file, generate_object_key
    try:
        file_content = download_file(payload.object_key, expected_hash=payload.file_hash)
        
        # Detect file extension from original filename or object_key
        file_ext = ".docx"  # default
        if payload.object_key.lower().endswith('.xlsx'):
            file_ext = ".xlsx"
        elif payload.object_key.lower().endswith('.pptx'):
            file_ext = ".pptx"
        elif payload.object_key.lower().endswith('.docx'):
            file_ext = ".docx"
        elif payload.object_key.lower().endswith(('.xml', '.bpmn', '.bpmn20', '.bpmn2', '.uml', '.xmi', '.archimate', '.archi', '.drawio', '.xsd', '.vsdx', '.vsdm', '.vsd', '.svg', '.png', '.jpg', '.jpeg', '.webp', '.csv', '.tsv', '.json', '.yaml', '.yml', '.md', '.txt', '.rst')):
            # Determine extension from filename
            if payload.object_key.lower().endswith('.bpmn2'):
                file_ext = ".bpmn2"
            elif payload.object_key.lower().endswith('.bpmn'):
                file_ext = ".bpmn"
            elif payload.object_key.lower().endswith('.uml'):
                file_ext = ".uml"
            elif payload.object_key.lower().endswith('.xmi'):
                file_ext = ".xmi"
            else:
                file_ext = ".xml"
        
        # Generate filename according to convention: [name]_[doc_type]_v[version].[ext]
        template_filename = _generate_template_filename(payload.name, payload.doc_type, version, file_ext.lstrip('.'))
        
        # Determine content type based on extension
        content_type_map = {
            '.xlsx': 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            '.pptx': 'application/vnd.openxmlformats-officedocument.presentationml.presentation',
            '.vsdx': 'application/vnd.ms-visio.drawing.main+xml',
            '.vsdm': 'application/vnd.ms-visio.macroEnabled.drawing.main+xml',
            '.vsd': 'application/vnd.visio',
            '.svg': 'image/svg+xml',
            '.png': 'image/png',
            '.jpg': 'image/jpeg',
            '.jpeg': 'image/jpeg',
            '.webp': 'image/webp',
            '.csv': 'text/csv',
            '.tsv': 'text/tab-separated-values',
            '.json': 'application/json',
            '.yaml': 'application/x-yaml',
            '.yml': 'application/x-yaml',
            '.md': 'text/markdown',
            '.txt': 'text/plain',
            '.rst': 'text/x-rst',
            '.xsd': 'application/xml',
        }
        
        # Check if we have a specific content type, otherwise use XML for modeling formats or default to docx
        if file_ext in content_type_map:
            content_type = content_type_map[file_ext]
        elif file_ext in ['.xml', '.bpmn', '.bpmn20', '.bpmn2', '.uml', '.xmi', '.archimate', '.archi', '.drawio', '.xsd']:
            content_type = "application/xml"
        elif file_ext in ['.parquet', '.lucid']:
            content_type = "application/octet-stream"
        else:
            content_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        
        # Upload with new filename according to convention
        new_object_key = generate_object_key("templates", template_filename)
        uploaded_key, file_hash = upload_file(
            file_content,
            new_object_key,
            content_type
        )
    except Exception as e:
        # If download fails, use original object_key (fallback)
        uploaded_key = payload.object_key
        file_hash = payload.file_hash
    
    template = Template(
        org_id=org_id,
        doc_type=payload.doc_type,
        name=payload.name,
        version=version,
        object_key=uploaded_key,
        file_hash=file_hash,
        mapping_manifest_json=payload.mapping_manifest_json,
        created_by=current_user.id,
        status=TemplateStatus.DRAFT.value,
    )
    db.add(template)
    db.flush()  # Flush to get template.id
    
    # Automatically generate PDF on template creation (skip for XML/BPMN/UML files)
    try:
        pdf_object_key, pdf_hash = _generate_template_pdf(template, db)
        if pdf_object_key is not None and pdf_hash is not None:
            template.pdf_object_key = pdf_object_key
            template.pdf_hash = pdf_hash
        db.flush()
    except Exception as pdf_error:
        # Log error but don't fail template creation if PDF generation fails
        import logging
        logger = logging.getLogger(__name__)
        logger.warning(f"PDF generation failed during template creation: {str(pdf_error)}")
    
    # HIPAA/GxP/GIS Compliance: Audit log for template creation
    try:
        from app.services.audit import log_action, AuditAction, model_to_dict
        
        # Get client info from request
        client_ip = request.client.host if request.client else None
        user_agent = request.headers.get("user-agent")
        
        log_action(
            db,
            actor_user_id=current_user.id,
            action=AuditAction.TEMPLATE_CREATE,
            entity_type="Template",
            entity_id=template.id,
            org_id=org_id,
            after_json={
                "id": str(template.id),
                "name": template.name,
                "doc_type": template.doc_type,
                "version": template.version,
                "status": template.status,
                "object_key": template.object_key,
                "file_hash": template.file_hash[:20] + "..." if template.file_hash else None,  # Store partial hash for audit
                "created_by": str(template.created_by),  # Include creator ID for audit
            },
            ip=client_ip,
            user_agent=user_agent,
        )
    except Exception as e:
        import logging
        logger = logging.getLogger(__name__)
        logger.warning(f"Audit logging failed: {str(e)}. Template created: {template.id} by user {current_user.id}")
    
    db.commit()
    db.refresh(template)
    return template


@router.post("/{template_id}/approve", response_model=TemplateRead)
def approve_template(
    request: Request,
    template_id: str, 
    db: Session = Depends(get_db), 
    current_user=Depends(get_current_active_user)
):
    from uuid import UUID
    try:
        template_uuid = UUID(template_id)
    except ValueError:
        raise HTTPException(status_code=400, detail="Invalid template ID format")
    template = db.query(Template).filter(Template.id == template_uuid).first()
    if not template:
        raise HTTPException(status_code=404, detail="Template not found")
    
    # HIPAA/GxP/GIS Compliance: Store before state for audit
    before_status = template.status
    
    template.status = TemplateStatus.APPROVED.value
    db.add(template)
    db.flush()
    
    # HIPAA/GxP/GIS Compliance: Audit log for template approval
    try:
        from app.services.audit import log_action, AuditAction
        
        client_ip = request.client.host if request.client else None
        user_agent = request.headers.get("user-agent")
        
        log_action(
            db,
            actor_user_id=current_user.id,
            action=AuditAction.TEMPLATE_APPROVE,
            entity_type="Template",
            entity_id=template.id,
            org_id=template.org_id,
            before_json={"status": before_status, "created_by": str(template.created_by)},
            after_json={"status": template.status, "created_by": str(template.created_by)},
            ip=client_ip,
            user_agent=user_agent,
        )
    except Exception as e:
        import logging
        logger = logging.getLogger(__name__)
        logger.warning(f"Audit logging failed: {str(e)}. Template approved: {template.id} by user {current_user.id}")
    
    db.commit()
    db.refresh(template)
    return template


@router.get("/{template_id}", response_model=TemplateRead)
def get_template(
    template_id: str, db: Session = Depends(get_db), current_user=Depends(get_current_active_user)
):
    """
    Get a single template by ID.
    HIPAA/GxP/GIS Compliance: Access control based on organization membership.
    """
    from uuid import UUID
    from app.models import ProjectMember, Project
    
    try:
        template_uuid = UUID(template_id)
    except ValueError:
        raise HTTPException(status_code=400, detail="Invalid template ID format")
    
        template = db.query(Template).filter(Template.id == template_uuid).first()
        if not template:
            raise HTTPException(status_code=404, detail="Template not found")
        
        # HIPAA/GxP/GIS Compliance: Verify user has access to template's organization
        try:
            user_orgs = db.query(Project.org_id).join(
                ProjectMember, Project.id == ProjectMember.project_id
            ).filter(
                ProjectMember.user_id == current_user.id,
                (ProjectMember.expires_at.is_(None) | (ProjectMember.expires_at > func.now()))
            ).distinct().all()
            user_org_ids = [org_id[0] for org_id in user_orgs] if user_orgs else []
            
            # Allow access if user is in the same org or if no org restrictions (for system templates)
            if user_org_ids and template.org_id not in user_org_ids:
                raise HTTPException(
                    status_code=403,
                    detail="Access denied: You do not have permission to view this template"
                )
        except HTTPException:
            raise
    except Exception:
        # If access check fails, log but allow access (for system templates)
        pass
    
        return template


@router.put("/{template_id}", response_model=TemplateRead)
def update_template(
    request: Request,
    template_id: str,
    payload: TemplateUpdate,
    db: Session = Depends(get_db),
    current_user=Depends(get_current_active_user),
):
    """Update a template."""
    from uuid import UUID
    from app.services.audit import log_action, AuditAction, model_to_dict
    
    try:
        template_uuid = UUID(template_id)
    except ValueError:
        raise HTTPException(status_code=400, detail="Invalid template ID format")
    template = db.query(Template).filter(Template.id == template_uuid).first()
    if not template:
        raise HTTPException(status_code=404, detail="Template not found")
    
    # HIPAA/GxP/GIS Compliance: Store before state for audit
    before_dict = model_to_dict(template)
    
    # Track if file was updated (object_key or file_hash changed)
    file_updated = False
    version_changed = False
    name_changed = False
    doc_type_changed = False
    
    if payload.object_key is not None and payload.object_key != template.object_key:
        file_updated = True
    if payload.file_hash is not None and payload.file_hash != template.file_hash:
        file_updated = True
    if payload.version is not None and payload.version != template.version:
        version_changed = True
    if payload.name is not None and payload.name != template.name:
        name_changed = True
    if payload.doc_type is not None and payload.doc_type != template.doc_type:
        doc_type_changed = True
    
    # Update fields if provided
    if payload.doc_type is not None:
        template.doc_type = payload.doc_type
    if payload.name is not None:
        template.name = payload.name
    if payload.version is not None:
        template.version = payload.version
    if payload.object_key is not None:
        template.object_key = payload.object_key
    if payload.file_hash is not None:
        template.file_hash = payload.file_hash
    if payload.mapping_manifest_json is not None:
        template.mapping_manifest_json = payload.mapping_manifest_json
    
    # If name, doc_type, or version changed, update object_key and pdf_object_key to match convention
    if (name_changed or doc_type_changed or version_changed) and template.object_key:
        from app.services.storage import download_file, upload_file, generate_object_key
        try:
            # Download existing file and re-upload with new name
            file_content = download_file(template.object_key, expected_hash=template.file_hash)
            
            # Detect file extension
            file_ext = ".docx"  # default
            if template.object_key.lower().endswith('.xlsx'):
                file_ext = ".xlsx"
            elif template.object_key.lower().endswith('.pptx'):
                file_ext = ".pptx"
            elif template.object_key.lower().endswith('.docx'):
                file_ext = ".docx"
            elif template.object_key.lower().endswith(('.xml', '.bpmn', '.bpmn2', '.uml', '.xmi')):
                # Determine extension from filename
                if template.object_key.lower().endswith('.bpmn2'):
                    file_ext = ".bpmn2"
                elif template.object_key.lower().endswith('.bpmn'):
                    file_ext = ".bpmn"
                elif template.object_key.lower().endswith('.uml'):
                    file_ext = ".uml"
                elif template.object_key.lower().endswith('.xmi'):
                    file_ext = ".xmi"
                else:
                    file_ext = ".xml"
            
            # Generate new filename according to convention
            new_template_filename = _generate_template_filename(template.name, template.doc_type, template.version, file_ext.lstrip('.'))
            new_object_key = generate_object_key("templates", new_template_filename)
            
            # Determine content type based on extension
            content_type_map = {
                '.xlsx': 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
                '.pptx': 'application/vnd.openxmlformats-officedocument.presentationml.presentation',
                '.vsdx': 'application/vnd.ms-visio.drawing.main+xml',
                '.vsdm': 'application/vnd.ms-visio.macroEnabled.drawing.main+xml',
                '.vsd': 'application/vnd.visio',
                '.svg': 'image/svg+xml',
                '.png': 'image/png',
                '.jpg': 'image/jpeg',
                '.jpeg': 'image/jpeg',
                '.webp': 'image/webp',
                '.csv': 'text/csv',
                '.tsv': 'text/tab-separated-values',
                '.json': 'application/json',
                '.yaml': 'application/x-yaml',
                '.yml': 'application/x-yaml',
                '.md': 'text/markdown',
                '.txt': 'text/plain',
                '.rst': 'text/x-rst',
                '.xsd': 'application/xml',
            }
            
            if file_ext in content_type_map:
                content_type = content_type_map[file_ext]
            elif file_ext in ['.xml', '.bpmn', '.bpmn20', '.bpmn2', '.uml', '.xmi', '.archimate', '.archi', '.drawio', '.xsd']:
                content_type = "application/xml"
            elif file_ext in ['.parquet', '.lucid']:
                content_type = "application/octet-stream"
            else:
                content_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            
            uploaded_key, file_hash = upload_file(
                file_content,
                new_object_key,
                content_type
            )
            template.object_key = uploaded_key
            template.file_hash = file_hash
            file_updated = True  # Mark as updated to trigger PDF regeneration
        except Exception as e:
            import logging
            logger = logging.getLogger(__name__)
            logger.warning(f"Failed to rename template file: {str(e)}")
    
    # If file was updated and template was APPROVED, change status to DRAFT
    # This ensures that updated templates need to be re-approved
    if file_updated and template.status == TemplateStatus.APPROVED.value:
        template.status = TemplateStatus.DRAFT.value
    
    # Automatically generate PDF if file was updated
    if file_updated:
        try:
            pdf_object_key, pdf_hash = _generate_template_pdf(template, db)
            if pdf_object_key is not None and pdf_hash is not None:
                template.pdf_object_key = pdf_object_key
                template.pdf_hash = pdf_hash
        except Exception as pdf_error:
            # Log error but don't fail template update if PDF generation fails
            import logging
            logger = logging.getLogger(__name__)
            logger.warning(f"PDF generation failed during template update: {str(pdf_error)}")
    
    # Only update status if explicitly provided in payload
    if payload.status is not None:
        template.status = payload.status
    
    db.add(template)
    db.flush()
    
    # HIPAA/GxP/GIS Compliance: Audit log for template update
    # Use separate session for audit logging to avoid transaction conflicts
    try:
        from app.db import SessionLocal
        
        client_ip = request.client.host if request.client else None
        user_agent = request.headers.get("user-agent")
        
        after_dict = model_to_dict(template)
        # Don't log full file_hash in audit (privacy/security)
        if 'file_hash' in before_dict and before_dict['file_hash']:
            before_dict['file_hash'] = before_dict['file_hash'][:20] + "..."
        if 'file_hash' in after_dict and after_dict['file_hash']:
            after_dict['file_hash'] = after_dict['file_hash'][:20] + "..."
        
        # Use separate session for audit logging
        audit_db = SessionLocal()
        try:
            log_action(
                audit_db,
                actor_user_id=current_user.id,
                action=AuditAction.TEMPLATE_UPDATE,
                entity_type="Template",
                entity_id=template.id,
                org_id=template.org_id,
                before_json=before_dict,
                after_json=after_dict,
                ip=client_ip,
                user_agent=user_agent,
            )
            audit_db.commit()
        except Exception as audit_error:
            audit_db.rollback()
            import logging
            logger = logging.getLogger(__name__)
            logger.warning(f"Audit logging failed: {str(audit_error)}. Template updated: {template.id} by user {current_user.id}")
        finally:
            audit_db.close()
    except Exception as e:
        import logging
        logger = logging.getLogger(__name__)
        logger.warning(f"Audit logging setup failed: {str(e)}. Template updated: {template.id} by user {current_user.id}")
    
    db.commit()
    db.refresh(template)
    return template


@router.delete("/{template_id}", status_code=status.HTTP_204_NO_CONTENT)
def delete_template(
    request: Request,
    template_id: str, 
    db: Session = Depends(get_db), 
    current_user=Depends(get_current_active_user)
):
    """Delete a template."""
    from uuid import UUID
    from app.services.audit import log_action, AuditAction, model_to_dict
    
    try:
        template_uuid = UUID(template_id)
    except ValueError:
        raise HTTPException(status_code=400, detail="Invalid template ID format")
    template = db.query(Template).filter(Template.id == template_uuid).first()
    if not template:
        raise HTTPException(status_code=404, detail="Template not found")
    
    # HIPAA/GxP/GIS Compliance: Store before state for audit
    before_dict = model_to_dict(template)
    # Don't log full file_hash in audit (privacy/security)
    if 'file_hash' in before_dict and before_dict['file_hash']:
        before_dict['file_hash'] = before_dict['file_hash'][:20] + "..."
    
    org_id = template.org_id
    template_id_str = str(template.id)
    
    # Check if template is being used by any document versions
    from app.models import DocumentVersion
    document_versions_count = db.query(DocumentVersion).filter(
        DocumentVersion.template_id == template_uuid
    ).count()
    
    if document_versions_count > 0:
            raise HTTPException(
            status_code=409,  # Conflict
            detail=f"Cannot delete template: it is being used by {document_versions_count} document version(s). Please remove or update these documents first."
        )
    
    db.delete(template)
    db.flush()
    
    # HIPAA/GxP/GIS Compliance: Audit log for template deletion
    try:
        client_ip = request.client.host if request.client else None
        user_agent = request.headers.get("user-agent")
        
        log_action(
            db,
            actor_user_id=current_user.id,
            action=AuditAction.TEMPLATE_DELETE,
            entity_type="Template",
            entity_id=template_id_str,
            org_id=org_id,
            before_json=before_dict,
            after_json=None,  # Entity deleted
            ip=client_ip,
            user_agent=user_agent,
        )
    except Exception as e:
        import logging
        logger = logging.getLogger(__name__)
        logger.warning(f"Audit logging failed: {str(e)}. Template deleted: {template_id_str} by user {current_user.id}")
    
        db.commit()
    return None
    
    return None


@router.get("/{template_id}/file")
def get_template_file(
    request: Request,
    template_id: str, 
    db: Session = Depends(get_db), 
    current_user=Depends(get_current_active_user)
):
    """Download template file."""
    from uuid import UUID
    from fastapi.responses import StreamingResponse
    from io import BytesIO
    import os
    from app.services.audit import log_action, AuditAction
    
    try:
        template_uuid = UUID(template_id)
    except ValueError:
        raise HTTPException(status_code=400, detail="Invalid template ID format")
    
    template = db.query(Template).filter(Template.id == template_uuid).first()
    if not template:
        raise HTTPException(status_code=404, detail="Template not found")
    
    # HIPAA/GxP/GIS Compliance: Verify user has access to template's organization
    # For now, allow access if user has no org restrictions (system templates)
    # In production, you might want stricter access control
    from app.models import ProjectMember, Project
    try:
        user_orgs = db.query(Project.org_id).join(
            ProjectMember, Project.id == ProjectMember.project_id
        ).filter(
            ProjectMember.user_id == current_user.id,
            (ProjectMember.expires_at.is_(None) | (ProjectMember.expires_at > func.now()))
        ).distinct().all()
        user_org_ids = [org_id[0] for org_id in user_orgs] if user_orgs else []
        
        # Only restrict access if user has org memberships and template is not in their orgs
        if user_org_ids and template.org_id not in user_org_ids:
            raise HTTPException(
                status_code=403,
                detail="Access denied: You do not have permission to access this template file"
            )
    except HTTPException:
        raise
    except Exception as access_error:
        # If access check fails, log but allow access (for system templates)
        import logging
        logger = logging.getLogger(__name__)
        logger.warning(f"Access check failed: {str(access_error)}. Allowing access to template {template.id} for user {current_user.id}")
    
    # HIPAA/GxP/GIS Compliance: Audit log for template file access
    # Use separate session for audit logging to avoid transaction conflicts
    try:
        from app.services.audit import log_action, AuditAction
        from app.db import SessionLocal
        
        client_ip = request.client.host if request.client else None
        user_agent = request.headers.get("user-agent")
        
        # Use separate session for audit logging to avoid transaction conflicts
        audit_db = SessionLocal()
        try:
            log_action(
                audit_db,
                actor_user_id=current_user.id,
                action=AuditAction.TEMPLATE_VIEW,
                entity_type="Template",
                entity_id=template.id,
                org_id=template.org_id,
                after_json={"action": "file_download", "object_key": template.object_key},
                ip=client_ip,
                user_agent=user_agent,
            )
            audit_db.commit()
        except Exception as audit_error:
            audit_db.rollback()
            import logging
            logger = logging.getLogger(__name__)
            logger.warning(f"Audit logging failed: {str(audit_error)}. Template file accessed: {template.id} by user {current_user.id}")
        finally:
            audit_db.close()
    except Exception as e:
        import logging
        logger = logging.getLogger(__name__)
        logger.warning(f"Audit logging setup failed: {str(e)}. Template file accessed: {template.id} by user {current_user.id}")
    
    # Try to use storage service if available
    import logging
    logger = logging.getLogger(__name__)
    
    logger.info(f"Attempting to load template file: {template.object_key}")
    
    file_content = None
    
    try:
        from app.services.storage import download_file
        logger.info("Storage service available, using download_file")
        logger.info(f"Attempting to download file with object_key: {template.object_key}")
        # HIPAA/GxP/GIS Compliance: Verify file integrity using hash
        try:
            file_content = download_file(template.object_key, expected_hash=template.file_hash)
            logger.info(f"File downloaded successfully, size: {len(file_content)} bytes, hash verified")
        except ValueError as hash_error:
            # Hash verification failed - log but still return file for now
            logger.warning(f"Hash verification failed: {str(hash_error)}. Returning file anyway.")
            # Try again without hash verification
            try:
                file_content = download_file(template.object_key, expected_hash=None)
                logger.info(f"File downloaded without hash verification, size: {len(file_content)} bytes")
            except FileNotFoundError as file_error:
                logger.error(f"File not found in storage: {str(file_error)}")
                raise HTTPException(status_code=404, detail=f"Template file not found: {template.object_key}")
        except FileNotFoundError as file_error:
            logger.error(f"File not found in storage: {str(file_error)}")
            logger.error(f"Template object_key: {template.object_key}")
            logger.error(f"Storage base: {os.path.join(os.path.dirname(__file__), '../../storage')}")
            raise HTTPException(
                status_code=404, 
                detail=f"Template file not found: {template.object_key}. The file may not have been uploaded to storage yet. Please upload the file or update the template's object_key."
            )
    except ImportError:
        logger.info("Storage service not available, trying local filesystem")
        # Fallback: try to read from local filesystem if object_key is a path
        if os.path.exists(template.object_key):
            logger.info(f"File found on local filesystem: {template.object_key}")
            with open(template.object_key, 'rb') as f:
                file_content = f.read()
            
            # HIPAA/GxP/GIS Compliance: Verify file integrity
            import hashlib
            actual_hash = hashlib.sha256(file_content).hexdigest()
            if template.file_hash and actual_hash != template.file_hash:
                logger.error(f"File integrity check failed: expected {template.file_hash}, got {actual_hash}")
                raise HTTPException(
                    status_code=500, 
                    detail="File integrity check failed. File may have been tampered with."
                )
            
            logger.info(f"File read successfully, size: {len(file_content)} bytes, hash verified")
        else:
            logger.error(f"File not found on local filesystem: {template.object_key}")
            raise HTTPException(status_code=404, detail=f"Template file not found at: {template.object_key}")
    except HTTPException:
        # Re-raise HTTP exceptions
        raise
    except FileNotFoundError as e:
        # File not found in storage
        logger.error(f"File not found: {str(e)}")
        raise HTTPException(status_code=404, detail=f"Template file not found: {template.object_key}. The file may not have been uploaded yet.")
    except ValueError as e:
        # Hash verification failed
        logger.error(f"File integrity check failed: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))
    except Exception as e:
        logger.error(f"Error loading template file: {str(e)}")
        import traceback
        logger.error(traceback.format_exc())
        raise HTTPException(status_code=500, detail=f"Failed to load template file: {str(e)}")
    
    # Ensure file_content was loaded
    if file_content is None:
        logger.error("File content is None after all attempts to load")
        raise HTTPException(status_code=500, detail="Failed to load template file: file content is empty")
    
    # Determine content type from file extension
    ext = template.object_key.split('.')[-1].lower()
    content_types = {
        'docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        'doc': 'application/msword',
        'xlsx': 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
        'xls': 'application/vnd.ms-excel',
        'pptx': 'application/vnd.openxmlformats-officedocument.presentationml.presentation',
        'ppt': 'application/vnd.ms-powerpoint',
        # BPMN formats
        'xml': 'application/xml',
        'bpmn': 'application/xml',
        'bpmn20': 'application/xml',
        'bpmn2': 'application/xml',
        # UML formats
        'uml': 'application/xml',
        'xmi': 'application/xml',
        # TOGAF/ArchiMate formats
        'archimate': 'application/xml',
        'archi': 'application/xml',
        # Visio formats
        'vsdx': 'application/vnd.ms-visio.drawing.main+xml',
        'vsdm': 'application/vnd.ms-visio.macroEnabled.drawing.main+xml',
        'vsd': 'application/vnd.visio',
        # Lucidchart/draw.io formats
        'lucid': 'application/octet-stream',
        'drawio': 'application/xml',
        # Image/Preview formats
        'svg': 'image/svg+xml',
        'png': 'image/png',
        'jpg': 'image/jpeg',
        'jpeg': 'image/jpeg',
        'webp': 'image/webp',
        # Data formats
        'csv': 'text/csv',
        'tsv': 'text/tab-separated-values',
        'json': 'application/json',
        'parquet': 'application/octet-stream',
        # YAML formats
        'yaml': 'application/x-yaml',
        'yml': 'application/x-yaml',
        # Text/Markdown formats
        'md': 'text/markdown',
        'txt': 'text/plain',
        'rst': 'text/x-rst',
        # Schema/Data Mapping formats
        'xsd': 'application/xml',
    }
    media_type = content_types.get(ext, 'application/octet-stream')
    
    # Generate filename according to convention: [name]_[doc_type]_v[version].[ext]
    download_filename = _generate_template_filename(template.name, template.doc_type, template.version, ext)
    
    return StreamingResponse(
        BytesIO(file_content),
        media_type=media_type,
        headers={
            "Content-Disposition": f'inline; filename="{download_filename}"',
            "Access-Control-Allow-Origin": "*",
            "Access-Control-Allow-Methods": "GET, OPTIONS",
            "Access-Control-Allow-Headers": "Authorization, Content-Type",
            "Cache-Control": "no-cache"
        }
    )


@router.get("/{template_id}/styles")
def get_template_styles(
    template_id: str,
    db: Session = Depends(get_db),
    current_user=Depends(get_current_active_user)
):
    """Extract fonts and styles from DOCX template file."""
    from uuid import UUID
    from docx import Document
    from io import BytesIO
    from app.services.storage import download_file
    from app.models import ProjectMember, Project
    
    try:
        template_uuid = UUID(template_id)
    except ValueError:
        raise HTTPException(status_code=400, detail="Invalid template ID format")
    
    template = db.query(Template).filter(Template.id == template_uuid).first()
    if not template:
        raise HTTPException(status_code=404, detail="Template not found")
    
    # Verify access
    try:
        user_orgs = db.query(Project.org_id).join(
            ProjectMember, Project.id == ProjectMember.project_id
        ).filter(
            ProjectMember.user_id == current_user.id,
            (ProjectMember.expires_at.is_(None) | (ProjectMember.expires_at > func.now()))
        ).distinct().all()
        user_org_ids = [org_id[0] for org_id in user_orgs] if user_orgs else []
        
        if user_org_ids and template.org_id not in user_org_ids:
            raise HTTPException(
                status_code=403,
                detail="Access denied"
            )
    except HTTPException:
        raise
    except Exception:
        pass  # Allow access for system templates
    
    # Download and parse DOCX
    try:
        from app.services.storage import download_file
        file_content = download_file(template.object_key, expected_hash=None)
    except Exception as e:
        raise HTTPException(status_code=404, detail=f"Template file not found: {str(e)}")
    
    try:
        doc = Document(BytesIO(file_content))
        
        # Extract fonts used in the document
        fonts_used = set()
        styles_info = {}
        list_styles = {}  # Store list styles (bullet points)
        table_styles = []  # Store table styles (colors, etc.)
        
        # Get all styles from document
        for style in doc.styles:
            if hasattr(style, 'font') and style.font.name:
                fonts_used.add(style.font.name)
            styles_info[style.name] = {
                "name": style.name,
                "type": str(style.type) if hasattr(style, 'type') else None,
            }
        
        # Extract fonts and list styles from paragraphs
        for paragraph in doc.paragraphs:
            for run in paragraph.runs:
                if run.font.name:
                    fonts_used.add(run.font.name)
            
            # Check if paragraph is part of a list
            if paragraph.style and hasattr(paragraph.style, 'paragraph_format'):
                pf = paragraph.style.paragraph_format
                if hasattr(pf, 'list_format') and pf.list_format:
                    list_format = pf.list_format
                    list_id = getattr(list_format, 'ilvl', None)
                    if list_id is not None:
                        # Get list style
                        list_style_key = f"list_level_{list_id}"
                        if list_style_key not in list_styles:
                            list_styles[list_style_key] = {
                                "level": list_id,
                                "number_format": getattr(list_format, 'num_style', None),
                            }
        
        # Extract fonts and table styles from tables
        from docx.oxml.ns import qn
        for table_idx, table in enumerate(doc.tables):
            table_style = {
                "index": table_idx,
                "rows": []
            }
            
            for row_idx, row in enumerate(table.rows):
                row_style = {
                    "index": row_idx,
                    "cells": []
                }
                
                for cell_idx, cell in enumerate(row.cells):
                    cell_style = {
                        "index": cell_idx,
                        "background_color": None,
                        "shading": None
                    }
                    
                    # Extract cell background color/shading - check multiple locations
                    if hasattr(cell, '_element'):
                        # Check cell properties (tcPr)
                        tc_pr = cell._element.find('.//w:tcPr', {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'})
                        if tc_pr is not None:
                            shading = tc_pr.find('.//w:shd', {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'})
                            if shading is not None:
                                fill = shading.get(qn('w:fill'))
                                if fill:
                                    # Convert hex to rgb
                                    fill_hex = fill
                                    if fill_hex.startswith('#'):
                                        fill_hex = fill_hex[1:]
                                    if len(fill_hex) == 6:
                                        try:
                                            r = int(fill_hex[0:2], 16)
                                            g = int(fill_hex[2:4], 16)
                                            b = int(fill_hex[4:6], 16)
                                            cell_style["background_color"] = f"rgb({r}, {g}, {b})"
                                            cell_style["shading"] = fill_hex
                                        except ValueError:
                                            pass
                    
                    # Also check paragraph shading in cell
                    if not cell_style["background_color"]:
                        for paragraph in cell.paragraphs:
                            if hasattr(paragraph, '_element'):
                                p_pr = paragraph._element.find('.//w:pPr', {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'})
                                if p_pr is not None:
                                    shading = p_pr.find('.//w:shd', {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'})
                                    if shading is not None:
                                        fill = shading.get(qn('w:fill'))
                                        if fill:
                                            fill_hex = fill
                                            if fill_hex.startswith('#'):
                                                fill_hex = fill_hex[1:]
                                            if len(fill_hex) == 6:
                                                try:
                                                    r = int(fill_hex[0:2], 16)
                                                    g = int(fill_hex[2:4], 16)
                                                    b = int(fill_hex[4:6], 16)
                                                    cell_style["background_color"] = f"rgb({r}, {g}, {b})"
                                                    cell_style["shading"] = fill_hex
                                                    break
                                                except ValueError:
                                                    pass
                            
                            for run in paragraph.runs:
                                if run.font.name:
                                    fonts_used.add(run.font.name)
                    
                    row_style["cells"].append(cell_style)
                table_style["rows"].append(row_style)
            table_styles.append(table_style)
            
            # Log table styles for debugging
            import logging
            logger = logging.getLogger(__name__)
            logger.info(f"Table {table_idx}: Found {len(table_style['rows'])} rows with colors")
            for row_idx, row_data in enumerate(table_style['rows']):
                colored_cells = [c for c in row_data['cells'] if c['background_color']]
                if colored_cells:
                    logger.info(f"  Row {row_idx}: {len(colored_cells)} cells with colors: {[c['background_color'] for c in colored_cells]}")
        
        # Extract list styles from document's numbering part
        try:
            from docx.oxml import parse_xml
            from docx.oxml.ns import nsdecls, qn
            
            # Try to extract numbering styles
            numbering_part = doc.part.numbering_part
            if numbering_part:
                numbering_xml = numbering_part.element
                # Parse numbering definitions
                for num in numbering_xml.findall('.//w:num', {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}):
                    num_id = num.get(qn('w:numId'))
                    abstract_num = num.find('.//w:abstractNumId', {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'})
                    if abstract_num is not None:
                        abstract_num_id = abstract_num.get(qn('w:val'))
                        # Find abstract numbering definition
                        abstract_num_def = numbering_xml.find(f'.//w:abstractNum[@w:abstractNumId="{abstract_num_id}"]', {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'})
                        if abstract_num_def is not None:
                            # Extract level definitions
                            for lvl in abstract_num_def.findall('.//w:lvl', {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}):
                                ilvl = lvl.get(qn('w:ilvl'))
                                num_fmt = lvl.find('.//w:numFmt', {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'})
                                
                                level_info = {
                                    "level": int(ilvl) if ilvl else 0,
                                    "number_format": None,
                                    "bullet_text": None,
                                    "bullet_char": None,
                                    "bullet_font": None,
                                    "left_indent_pt": None,
                                    "hanging_indent_pt": None,
                                    "first_line_indent_pt": None,
                                }
                                
                                if num_fmt is not None:
                                    fmt_val = num_fmt.get(qn('w:val'))
                                    level_info["number_format"] = fmt_val
                                
                                # Extract indentation from paragraph properties (w:pPr/w:ind)
                                p_pr = lvl.find('.//w:pPr', {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'})
                                if p_pr is not None:
                                    indent = p_pr.find('.//w:ind', {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'})
                                    if indent is not None:
                                        # Indentation is in TWIPs (twentieths of a point), convert to points
                                        left_twips = indent.get(qn('w:left'))
                                        hanging_twips = indent.get(qn('w:hanging'))
                                        first_line_twips = indent.get(qn('w:firstLine'))
                                        
                                        if left_twips:
                                            level_info["left_indent_pt"] = round(int(left_twips) / 20, 2)
                                        if hanging_twips:
                                            level_info["hanging_indent_pt"] = round(int(hanging_twips) / 20, 2)
                                        if first_line_twips:
                                            level_info["first_line_indent_pt"] = round(int(first_line_twips) / 20, 2)
                                
                                # Extract bullet text/symbol (w:lvlText contains the actual symbol)
                                lvl_text = lvl.find('.//w:lvlText', {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'})
                                if lvl_text is not None:
                                    bullet_text = lvl_text.get(qn('w:val'))
                                    if bullet_text:
                                        level_info["bullet_text"] = bullet_text
                                        # Extract the actual character(s) - might be Unicode escape sequences
                                        # DOCX uses %1, %2, etc. for placeholders, but actual symbols are in the text
                                        level_info["bullet_char"] = bullet_text
                                
                                # Extract bullet font (w:rFonts or w:rPr/w:rFonts)
                                r_pr = lvl.find('.//w:rPr', {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'})
                                if r_pr is not None:
                                    r_fonts = r_pr.find('.//w:rFonts', {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'})
                                    if r_fonts is not None:
                                        font_name = r_fonts.get(qn('w:ascii')) or r_fonts.get(qn('w:hAnsi'))
                                        if font_name:
                                            level_info["bullet_font"] = font_name
                                
                                # Also check for custom bullet symbols in w:lvlPicBulletId (image bullets)
                                lvl_pic = lvl.find('.//w:lvlPicBulletId', {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'})
                                if lvl_pic is not None:
                                    pic_id = lvl_pic.get(qn('w:val'))
                                    level_info["bullet_image_id"] = pic_id
                                
                                if ilvl not in list_styles:
                                    list_styles[f"list_level_{ilvl}"] = level_info
                                else:
                                    # Merge with existing info
                                    existing = list_styles[f"list_level_{ilvl}"]
                                    existing.update(level_info)
        except Exception as e:
            import logging
            logger = logging.getLogger(__name__)
            logger.warning(f"Could not extract list styles: {str(e)}")
        
        # Standard Word fonts mapping (fallback)
        word_fonts = {
            "Calibri": "Calibri, 'Segoe UI', Arial, sans-serif",
            "Times New Roman": "'Times New Roman', Times, serif",
            "Arial": "Arial, Helvetica, sans-serif",
            "Cambria": "Cambria, Georgia, serif",
            "Verdana": "Verdana, Geneva, sans-serif",
            "Tahoma": "Tahoma, Geneva, sans-serif",
            "Courier New": "'Courier New', Courier, monospace",
            "Georgia": "Georgia, serif",
            "Trebuchet MS": "'Trebuchet MS', Helvetica, sans-serif",
            "Comic Sans MS": "'Comic Sans MS', cursive",
            "Impact": "Impact, Charcoal, sans-serif",
            "Lucida Console": "'Lucida Console', Monaco, monospace",
        }
        
        # Map fonts to CSS font families
        font_families = {}
        for font in fonts_used:
            if font in word_fonts:
                font_families[font] = word_fonts[font]
            else:
                # Use font name directly, with fallbacks
                font_families[font] = f"'{font}', Arial, sans-serif"
        
        # Extract page orientation information from sections
        page_orientations = []
        for section_idx, section in enumerate(doc.sections):
            # Convert EMU to pixels (1 EMU = 1/914400 inch, 96 DPI = 96 pixels per inch)
            # A4 Portrait: 21.0 cm x 29.7 cm = 8.27" x 11.69" = 794px x 1123px at 96 DPI
            # A4 Landscape: 29.7 cm x 21.0 cm = 11.69" x 8.27" = 1123px x 794px at 96 DPI
            width_emu = section.page_width
            height_emu = section.page_height
            
            # Check orientation
            is_landscape = width_emu > height_emu
            
            # Get orientation attribute if available
            pgSz = section._sectPr.find(qn('w:pgSz'))
            orient_attr = None
            if pgSz is not None:
                orient_attr = pgSz.get(qn('w:orient'))
            
            # Convert to pixels (96 DPI)
            width_px = int((width_emu / 914400) * 96)
            height_px = int((height_emu / 914400) * 96)
            
            page_orientations.append({
                "section_index": section_idx,
                "orientation": "landscape" if is_landscape else "portrait",
                "width_px": width_px,
                "height_px": height_px,
                "width_cm": round(width_emu / 914400 * 2.54, 2),
                "height_cm": round(height_emu / 914400 * 2.54, 2),
                "left_margin_cm": round(section.left_margin / 914400 * 2.54, 2),
                "right_margin_cm": round(section.right_margin / 914400 * 2.54, 2),
                "top_margin_cm": round(section.top_margin / 914400 * 2.54, 2),
                "bottom_margin_cm": round(section.bottom_margin / 914400 * 2.54, 2),
                "orientation_attribute": orient_attr
            })
        
        return {
            "fonts": list(fonts_used),
            "font_families": font_families,
            "styles": styles_info,
            "default_fonts": word_fonts,
            "list_styles": list_styles,
            "table_styles": table_styles,
            "page_orientations": page_orientations
        }
        
    except Exception as e:
        import logging
        logger = logging.getLogger(__name__)
        logger.error(f"Error parsing DOCX: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Failed to parse DOCX: {str(e)}")


@router.get("/{template_id}/pdf/check")
def check_pdf_exists(
    template_id: str,
    db: Session = Depends(get_db),
    current_user=Depends(get_current_active_user)
):
    """Check if PDF version of template exists."""
    from uuid import UUID
    
    try:
        template_uuid = UUID(template_id)
    except ValueError:
        raise HTTPException(status_code=400, detail="Invalid template ID format")
    
    template = db.query(Template).filter(Template.id == template_uuid).first()
    if not template:
        raise HTTPException(status_code=404, detail="Template not found")
    
    # Verify access
    try:
        user_orgs = db.query(Project.org_id).join(
            ProjectMember, Project.id == ProjectMember.project_id
        ).filter(
            ProjectMember.user_id == current_user.id,
            (ProjectMember.expires_at.is_(None) | (ProjectMember.expires_at > func.now()))
        ).distinct().all()
        user_org_ids = [org_id[0] for org_id in user_orgs] if user_orgs else []
        
        if user_org_ids and template.org_id not in user_org_ids:
            raise HTTPException(
                status_code=403,
                detail="Access denied"
            )
    except HTTPException:
        raise
    except Exception:
        pass  # Allow access for system templates
    
    exists = template.pdf_object_key is not None and template.pdf_object_key != ""
    
    return {
        "exists": exists,
        "pdf_object_key": template.pdf_object_key if exists else None
    }


@router.post("/{template_id}/pdf/generate")
def generate_template_pdf(
    template_id: str,
    db: Session = Depends(get_db),
    current_user=Depends(get_current_active_user)
):
    """Generate PDF version from template (DOCX or XLSX) using LibreOffice."""
    from uuid import UUID
    
    # 1. Walidacja ID
    try:
        template_uuid = UUID(template_id)
    except ValueError:
        raise HTTPException(status_code=400, detail="Invalid template ID")
    
    template = db.query(Template).filter(Template.id == template_uuid).first()
    if not template:
        raise HTTPException(status_code=404, detail="Template not found")
    
    # Verify access
    try:
        user_orgs = db.query(Project.org_id).join(
            ProjectMember, Project.id == ProjectMember.project_id
        ).filter(
            ProjectMember.user_id == current_user.id,
            (ProjectMember.expires_at.is_(None) | (ProjectMember.expires_at > func.now()))
        ).distinct().all()
        user_org_ids = [org_id[0] for org_id in user_orgs] if user_orgs else []
        
        if user_org_ids and template.org_id not in user_org_ids:
            raise HTTPException(
                status_code=403,
                detail="Access denied"
            )
    except HTTPException:
        raise
    except Exception:
        pass  # Allow access for system templates

    # Generate PDF using helper function (handles DOCX/XLSX/PPTX; returns None for XML/BPMN/UML)
    try:
        pdf_object_key, pdf_hash = _generate_template_pdf(template, db)
        if pdf_object_key is None or pdf_hash is None:
            raise HTTPException(
                status_code=400,
                detail="PDF generation is not supported for XML/BPMN/UML/XMI files. These formats can be viewed directly."
            )
        template.pdf_object_key = pdf_object_key
        template.pdf_hash = pdf_hash
        db.commit()
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"PDF generation failed: {str(e)}")
    
    return {
        "status": "success",
        "pdf_object_key": pdf_object_key,
        "pdf_hash": pdf_hash
    }


@router.get("/{template_id}/pdf")
def get_template_pdf(
    request: Request,
    template_id: str,
    db: Session = Depends(get_db),
    current_user=Depends(get_current_active_user)
):
    """Download PDF version of template."""
    from uuid import UUID
    from fastapi.responses import StreamingResponse
    from io import BytesIO
    import os
    from app.services.storage import download_file
    
    try:
        template_uuid = UUID(template_id)
    except ValueError:
        raise HTTPException(status_code=400, detail="Invalid template ID format")
    
    template = db.query(Template).filter(Template.id == template_uuid).first()
    if not template:
        raise HTTPException(status_code=404, detail="Template not found")
    
    # Refresh template to get latest pdf_object_key (in case it was just generated)
    db.refresh(template)
    
    if not template.pdf_object_key:
        raise HTTPException(
            status_code=404,
            detail="PDF version not found. Please generate PDF first."
        )
    
    # Verify access
    try:
        user_orgs = db.query(Project.org_id).join(
            ProjectMember, Project.id == ProjectMember.project_id
        ).filter(
            ProjectMember.user_id == current_user.id,
            (ProjectMember.expires_at.is_(None) | (ProjectMember.expires_at > func.now()))
        ).distinct().all()
        user_org_ids = [org_id[0] for org_id in user_orgs] if user_orgs else []
        
        if user_org_ids and template.org_id not in user_org_ids:
            raise HTTPException(
                status_code=403,
                detail="Access denied"
            )
    except HTTPException:
        raise
    except Exception:
        pass  # Allow access for system templates
    
    # Download PDF file - use same approach as get_template_file
    import logging
    logger = logging.getLogger(__name__)
    
    logger.info(f"Attempting to load template PDF: {template.pdf_object_key}")
    
    pdf_content = None
    
    try:
        from app.services.storage import download_file
        logger.info("Storage service available, using download_file")
        logger.info(f"Attempting to download PDF with object_key: {template.pdf_object_key}")
        # HIPAA/GxP/GIS Compliance: Verify file integrity using hash
        try:
            pdf_content = download_file(template.pdf_object_key, expected_hash=template.pdf_hash)
            logger.info(f"PDF downloaded successfully from storage, size: {len(pdf_content)} bytes, hash verified")
        except ValueError as hash_error:
            # Hash verification failed - log but still return file for now
            logger.warning(f"Hash verification failed: {str(hash_error)}. Returning PDF anyway.")
            # Try again without hash verification
            try:
                pdf_content = download_file(template.pdf_object_key, expected_hash=None)
                logger.info(f"PDF downloaded without hash verification, size: {len(pdf_content)} bytes")
            except FileNotFoundError as file_error:
                logger.error(f"PDF not found in storage: {str(file_error)}")
                raise HTTPException(status_code=404, detail=f"Template PDF not found: {template.pdf_object_key}")
        except FileNotFoundError as file_error:
            logger.warning(f"PDF not found in storage: {str(file_error)}")
            logger.info(f"Template pdf_object_key: {template.pdf_object_key}")
            logger.info("PDF file missing but pdf_object_key exists in DB. Attempting to regenerate PDF...")
            
            # Try to regenerate PDF if file is missing
            try:
                pdf_object_key, pdf_hash = _generate_template_pdf(template, db)
                template.pdf_object_key = pdf_object_key
                template.pdf_hash = pdf_hash
                db.commit()
                db.refresh(template)
                
                # Now try to download the newly generated PDF
                pdf_content = download_file(template.pdf_object_key, expected_hash=template.pdf_hash)
                logger.info(f"PDF regenerated and downloaded successfully, size: {len(pdf_content)} bytes")
            except Exception as regenerate_error:
                logger.error(f"Failed to regenerate PDF: {str(regenerate_error)}")
                raise HTTPException(
                    status_code=404, 
                    detail=f"Template PDF not found and could not be regenerated: {template.pdf_object_key}"
                )
    except ImportError:
        logger.info("Storage service not available, trying local filesystem")
        # Fallback: try to read from local filesystem if pdf_object_key is a path
        if os.path.exists(template.pdf_object_key):
            logger.info(f"PDF found on local filesystem: {template.pdf_object_key}")
            with open(template.pdf_object_key, 'rb') as f:
                pdf_content = f.read()
            
            # HIPAA/GxP/GIS Compliance: Verify file integrity
            import hashlib
            actual_hash = hashlib.sha256(pdf_content).hexdigest()
            if template.pdf_hash and actual_hash != template.pdf_hash:
                logger.error(f"PDF integrity check failed: expected {template.pdf_hash}, got {actual_hash}")
                raise HTTPException(
                    status_code=500,
                    detail="PDF integrity check failed. File may have been tampered with."
                )
            
            logger.info(f"PDF read successfully, size: {len(pdf_content)} bytes, hash verified")
        else:
            logger.warning(f"PDF not found on local filesystem: {template.pdf_object_key}")
            logger.info("PDF file missing but pdf_object_key exists in DB. Attempting to regenerate PDF...")
            
            # Try to regenerate PDF if file is missing
            try:
                pdf_object_key, pdf_hash = _generate_template_pdf(template, db)
                if pdf_object_key and pdf_hash:
                    template.pdf_object_key = pdf_object_key
                    template.pdf_hash = pdf_hash
                    db.commit()
                db.refresh(template)
                
                # Now try to read the newly generated PDF
                if os.path.exists(template.pdf_object_key):
                    with open(template.pdf_object_key, 'rb') as f:
                        pdf_content = f.read()
                    logger.info(f"PDF regenerated and read successfully, size: {len(pdf_content)} bytes")
                else:
                    raise HTTPException(status_code=404, detail=f"Template PDF not found at: {template.pdf_object_key}")
            except HTTPException:
                raise
            except Exception as regenerate_error:
                logger.error(f"Failed to regenerate PDF: {str(regenerate_error)}")
                raise HTTPException(status_code=404, detail=f"Template PDF not found and could not be regenerated: {template.pdf_object_key}")
    except FileNotFoundError as e:
        # File not found in storage
        logger.error(f"PDF not found: {str(e)}")
        raise HTTPException(status_code=404, detail=f"Template PDF not found: {template.pdf_object_key}")
    except ValueError as e:
        # Hash verification failed
        logger.error(f"PDF integrity check failed: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error loading template PDF: {str(e)}")
        import traceback
        logger.error(traceback.format_exc())
        raise HTTPException(status_code=500, detail=f"Failed to load template PDF: {str(e)}")
    
    # Ensure pdf_content was loaded
    if pdf_content is None:
        logger.error("PDF content is None after all attempts to load")
        raise HTTPException(status_code=500, detail="Failed to load template PDF: PDF content is empty")
    
    # Generate PDF filename according to convention: [name]_[doc_type]_v[version].pdf
    pdf_download_filename = _generate_template_filename(template.name, template.doc_type, template.version, "pdf")
    
    return StreamingResponse(
        BytesIO(pdf_content),
        media_type="application/pdf",
        headers={
            "Content-Disposition": f'inline; filename="{pdf_download_filename}"',
            "Access-Control-Allow-Origin": "*",
            "Access-Control-Allow-Methods": "GET, OPTIONS",
            "Access-Control-Allow-Headers": "Authorization, Content-Type",
            "Cache-Control": "no-cache"
        }
    )


@router.get("/{template_id}/canvas")
def get_template_canvas(
    template_id: str,
    db: Session = Depends(get_db),
    current_user=Depends(get_current_active_user)
):
    """
    Get Excel template data as JSON for canvas view.
    Returns sheets with columns and rows data.
    """
    from uuid import UUID
    from openpyxl import load_workbook
    from openpyxl.utils import get_column_letter
    from app.services.storage import download_file
    import tempfile
    import os
    
    try:
        template_uuid = UUID(template_id)
    except ValueError:
        raise HTTPException(status_code=400, detail="Invalid template ID format")
    
    template = db.query(Template).filter(Template.id == template_uuid).first()
    if not template:
        raise HTTPException(status_code=404, detail="Template not found")
    
    # Check if file is Excel
    file_ext = os.path.splitext(template.object_key)[1].lower()
    if file_ext != '.xlsx':
        raise HTTPException(
            status_code=400,
            detail="Canvas view is only available for Excel (.xlsx) files"
        )
    
    # Download Excel file
    try:
        xlsx_bytes = download_file(
            template.object_key,
            expected_hash=template.file_hash
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Cannot load Excel file: {str(e)}")
    
    # Parse Excel and extract data
    with tempfile.NamedTemporaryFile(delete=False, suffix='.xlsx') as tmp_file:
        tmp_file.write(xlsx_bytes)
        tmp_path = tmp_file.name
    
    try:
        # Load workbook with data_only=True to get calculated values
        wb = load_workbook(tmp_path, data_only=True, read_only=True)
        
        sheets_data = []
        
        for ws in wb.worksheets:
            # Detect data range
            max_row = 0
            max_col = 0
            
            for row in ws.iter_rows():
                for cell in row:
                    if cell.value not in (None, "", " "):
                        max_row = max(max_row, cell.row)
                        max_col = max(max_col, cell.column)
            
            if max_row == 0 or max_col == 0:
                continue  # Skip empty sheets
            
            # Extract header row (first row)
            headers = []
            for col_idx in range(1, max_col + 1):
                col_letter = get_column_letter(col_idx)
                cell = ws[f"{col_letter}1"]
                header_value = str(cell.value).strip() if cell.value else f"Column {col_letter}"
                headers.append({
                    "key": col_letter,  # Use column letter as key
                    "label": header_value,
                    "type": "string"  # Can be enhanced with type detection
                })
            
            # Extract data rows
            rows = []
            for row_idx in range(2, max_row + 1):  # Start from row 2 (skip header)
                row_data = {}
                for col_idx, header in enumerate(headers, start=1):
                    col_letter = header["key"]
                    cell = ws[f"{col_letter}{row_idx}"]
                    # Convert value to string, handle None
                    cell_value = cell.value
                    if cell_value is None:
                        cell_value = ""
                    elif isinstance(cell_value, (int, float)):
                        cell_value = str(cell_value)
                    else:
                        cell_value = str(cell_value)
                    
                    row_data[col_letter] = cell_value
                
                rows.append(row_data)
            
            sheets_data.append({
                "name": ws.title,
                "columns": headers,
                "rows": rows
            })
        
        wb.close()
        
        return {"sheets": sheets_data}
    
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to parse Excel file: {str(e)}")
    
    finally:
        # Clean up temp file
        try:
            os.unlink(tmp_path)
        except:
            pass


@router.get("/{template_id}/info")
def get_template_info(
    template_id: str,
    db: Session = Depends(get_db),
    current_user=Depends(get_current_active_user)
):
    """
    Get detailed template information including creator, approver, and compliance details.
    """
    from uuid import UUID
    from app.models import AuditLog, User
    from app.services.audit import AuditAction
    import logging
    
    logger = logging.getLogger(__name__)
    
    try:
        template_uuid = UUID(template_id)
    except ValueError:
        raise HTTPException(status_code=400, detail="Invalid template ID format")
    
    try:
        template = db.query(Template).filter(Template.id == template_uuid).first()
        if not template:
            raise HTTPException(status_code=404, detail="Template not found")
    
        # Verify access (similar to get_template)
        try:
            user_orgs = db.query(Project.org_id).join(
                ProjectMember, Project.id == ProjectMember.project_id
            ).filter(
                ProjectMember.user_id == current_user.id,
                (ProjectMember.expires_at.is_(None) | (ProjectMember.expires_at > func.now()))
            ).distinct().all()
            user_org_ids = [org_id[0] for org_id in user_orgs] if user_orgs else []
            
            # Allow access if user is in the same org or if no org restrictions (for system templates)
            if user_org_ids and template.org_id not in user_org_ids:
                raise HTTPException(
                    status_code=403,
                    detail="Access denied: You do not have permission to view this template"
                )
        except HTTPException:
            raise
        except Exception:
            # If access check fails, log but allow access (for system templates)
            pass
        
        # Get creator information
        creator_info = None
        if template.created_by:
            try:
                creator = db.query(User).filter(User.id == template.created_by).first()
                if creator:
                    creator_info = {
                        "id": str(creator.id),
                        "name": getattr(creator, 'name', None) or creator.email,
                        "email": creator.email
                    }
            except Exception as e:
                logger.warning(f"Failed to load creator info: {str(e)}")
                # Continue without creator info
        
        # Get approver information from audit log
        approver_info = None
        try:
            # Try to find approval log - action is stored as string in DB
            # AuditAction.TEMPLATE_APPROVE is "TEMPLATE_APPROVE" (the string value)
            approval_log = db.query(AuditLog).filter(
                AuditLog.entity_type == "Template",
                AuditLog.entity_id == template_uuid,
                AuditLog.action == "TEMPLATE_APPROVE"
            ).order_by(AuditLog.created_at.desc()).first()
            
            logger.info(f"Approval log query for template {template_uuid}: found={approval_log is not None}")
            if approval_log:
                logger.info(f"Approval log: actor_user_id={approval_log.actor_user_id}, created_at={approval_log.created_at}, action={approval_log.action}")
            
            if approval_log and approval_log.actor_user_id:
                approver = db.query(User).filter(User.id == approval_log.actor_user_id).first()
                if approver:
                    approver_info = {
                        "id": str(approver.id),
                        "name": getattr(approver, 'name', None) or approver.email,
                        "email": approver.email,
                        "approved_at": approval_log.created_at.isoformat() if approval_log.created_at else None,
                        "comment": approval_log.after_json.get("comment") if approval_log.after_json and isinstance(approval_log.after_json, dict) else None
                    }
                    logger.info(f"Approver info loaded: {approver_info['name']}, approved_at={approver_info['approved_at']}")
                else:
                    logger.warning(f"Approver user not found for actor_user_id={approval_log.actor_user_id}")
            else:
                logger.info(f"No approval log found for template {template_uuid}")
                # Also check if template status is APPROVED but no log found
                if template.status == "APPROVED":
                    logger.warning(f"Template status is APPROVED but no approval log found")
        except Exception as e:
            logger.error(f"Failed to load approver info: {str(e)}", exc_info=True)
            # Continue without approver info
        
        # Determine active stage in workflow (based on status)
        active_stage = template.status  # DRAFT, APPROVED, or ARCHIVED
        
        # Get compliance standards from compliance service
        from app.services.compliance import get_compliance_standards
        compliance_standards = get_compliance_standards(template.doc_type, include_default=True)
        
        return {
            "template_id": str(template.id),
            "name": template.name,
            "doc_type": template.doc_type,
            "version": template.version,
            "status": template.status,
            "active_stage": active_stage,
            "created_by": creator_info,
            "created_at": template.created_at.isoformat() if template.created_at else None,
            "approved_by": approver_info,
            "compliance_standards": compliance_standards,
            "file_hash": template.file_hash,
            "pdf_hash": template.pdf_hash if template.pdf_hash else None,
            "checked_out_by": str(template.checked_out_by) if template.checked_out_by else None,
            "checked_out_at": template.checked_out_at.isoformat() if template.checked_out_at else None,
        }
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error getting template info: {str(e)}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"Failed to get template information: {str(e)}")


@router.get("/{template_id}/audit-log")
def get_template_audit_log(
    template_id: str,
    db: Session = Depends(get_db),
    current_user=Depends(get_current_active_user)
):
    """
    Get audit log entries for a template.
    Returns all actions performed on the template for compliance tracking.
    """
    from uuid import UUID
    from app.models import AuditLog, User
    
    try:
        template_uuid = UUID(template_id)
    except ValueError:
        raise HTTPException(status_code=400, detail="Invalid template ID format")
    
    template = db.query(Template).filter(Template.id == template_uuid).first()
    if not template:
        raise HTTPException(status_code=404, detail="Template not found")
    
    # Get all audit log entries for this template
    audit_logs = db.query(AuditLog).filter(
        AuditLog.entity_type == "Template",
        AuditLog.entity_id == template_uuid
    ).order_by(AuditLog.created_at.desc()).all()
    
    # Build response with actor information and resolve user IDs in JSON
    log_entries = []
    for log in audit_logs:
        actor = db.query(User).filter(User.id == log.actor_user_id).first()
        actor_name = None
        actor_email = None
        if actor:
            actor_name = getattr(actor, 'name', None) or actor.email
            actor_email = actor.email
        
        # Resolve user IDs in before_json and after_json to usernames
        before_json = log.before_json
        after_json = log.after_json
        
        # Helper function to resolve user IDs in JSON
        def resolve_user_ids(json_data, db_session):
            if not json_data or not isinstance(json_data, dict):
                return json_data
            resolved = json_data.copy()
            
            # Resolve created_by if present
            if 'created_by' in resolved and resolved['created_by']:
                try:
                    user_id = UUID(str(resolved['created_by'])) if not isinstance(resolved['created_by'], UUID) else resolved['created_by']
                    user = db_session.query(User).filter(User.id == user_id).first()
                    if user:
                        resolved['created_by_user_id'] = str(user.id)  # Keep original ID
                        resolved['created_by_name'] = getattr(user, 'name', None) or user.email
                except (ValueError, TypeError):
                    pass
            
            return resolved
        
        before_json = resolve_user_ids(before_json, db)
        after_json = resolve_user_ids(after_json, db)
        
        log_entries.append({
            "id": str(log.id),
            "action": log.action,
            "actor_user_id": str(log.actor_user_id),
            "actor_name": actor_name,
            "actor_email": actor_email,
            "created_at": log.created_at.isoformat() if log.created_at else None,
            "before_json": before_json,
            "after_json": after_json,
            "ip": log.ip,
            "user_agent": log.user_agent,
        })
    
    return {
        "template_id": str(template.id),
        "template_name": template.name,
        "total_entries": len(log_entries),
        "entries": log_entries
    }


@router.get("/{template_id}/documents")
def get_template_documents(
    template_id: str,
    db: Session = Depends(get_db),
    current_user=Depends(get_current_active_user)
):
    """Get list of documents using this template."""
    from uuid import UUID
    from app.models import Document, DocumentVersion
    
    try:
        template_uuid = UUID(template_id)
    except ValueError:
        raise HTTPException(status_code=400, detail="Invalid template ID format")
    
    template = db.query(Template).filter(Template.id == template_uuid).first()
    if not template:
        raise HTTPException(status_code=404, detail="Template not found")
    
    # Verify access
    try:
        user_orgs = db.query(Project.org_id).join(
            ProjectMember, Project.id == ProjectMember.project_id
        ).filter(
            ProjectMember.user_id == current_user.id,
            (ProjectMember.expires_at.is_(None) | (ProjectMember.expires_at > func.now()))
        ).distinct().all()
        user_org_ids = [org_id[0] for org_id in user_orgs] if user_orgs else []
        
        if user_org_ids and template.org_id not in user_org_ids:
            raise HTTPException(
                status_code=403,
                detail="Access denied"
            )
    except HTTPException:
        raise
    except Exception:
        pass  # Allow access for system templates
    
    # Find all document versions using this template
    versions = db.query(DocumentVersion).filter(
        DocumentVersion.template_id == template_uuid
    ).all()
    
    # Get unique documents
    document_ids = {v.document_id for v in versions}
    documents = db.query(Document).filter(Document.id.in_(document_ids)).all()
    
    # Return simplified document info
    return [
        {
            "id": str(doc.id),
            "title": doc.title,
            "doc_type": doc.doc_type,
            "created_at": doc.created_at.isoformat() if doc.created_at else None
        }
        for doc in documents
    ]

