"""
DOCX Validation Service for PDF Conversion Compatibility

Validates DOCX files to ensure they will convert properly to PDF using LibreOffice.
Based on enterprise-grade best practices for DOCX → PDF conversion.
"""
import logging
from typing import Dict, List, Optional, Tuple
from io import BytesIO
from docx import Document
from docx.oxml.ns import qn

logger = logging.getLogger(__name__)


class DOCXValidationResult:
    """Result of DOCX validation with warnings and errors."""
    
    def __init__(self):
        self.is_valid = True
        self.warnings: List[Dict[str, str]] = []
        self.errors: List[Dict[str, str]] = []
        self.summary: Dict[str, int] = {
            'images_embedded': 0,
            'images_linked': 0,
            'images_inline': 0,
            'images_floating': 0,
            'theme_colors': 0,
            'rgb_colors': 0,
            'text_boxes': 0,
            'smart_art': 0,
            'toc_fields': 0,
            'cover_pages': 0,
        }
    
    def add_warning(self, category: str, message: str, fix_hint: str = ""):
        """Add a warning (non-blocking issue)."""
        self.warnings.append({
            'category': category,
            'message': message,
            'fix_hint': fix_hint
        })
    
    def add_error(self, category: str, message: str, fix_hint: str = ""):
        """Add an error (blocking issue)."""
        self.is_valid = False
        self.errors.append({
            'category': category,
            'message': message,
            'fix_hint': fix_hint
        })


def validate_docx_for_pdf_conversion(docx_content: bytes) -> DOCXValidationResult:
    """
    Validate DOCX file for PDF conversion compatibility with LibreOffice.
    
    Checks:
    1. Images are embedded (not linked)
    2. Images are inline (not floating)
    3. Colors are RGB (not theme colors)
    4. No TextBoxes
    5. No SmartArt
    6. TOC is updated (if present)
    7. No Word Cover Page
    
    Returns:
        DOCXValidationResult with warnings and errors
    """
    result = DOCXValidationResult()
    
    try:
        doc = Document(BytesIO(docx_content))
        
        # 1. Check images (embedded vs linked, inline vs floating)
        _validate_images(doc, result)
        
        # 2. Check colors (RGB vs theme colors)
        _validate_colors(doc, result)
        
        # 3. Check for TextBoxes
        _validate_text_boxes(doc, result)
        
        # 4. Check for SmartArt
        _validate_smart_art(doc, result)
        
        # 5. Check TOC fields
        _validate_toc(doc, result)
        
        # 6. Check for Cover Page
        _validate_cover_page(doc, result)
        
        # 7. Check fonts (warn if system fonts might not be available)
        _validate_fonts(doc, result)
        
    except Exception as e:
        logger.error(f"Error validating DOCX: {str(e)}", exc_info=True)
        result.add_error(
            'parsing',
            f"Failed to parse DOCX file: {str(e)}",
            "Ensure the file is a valid DOCX document"
        )
    
    return result


def _validate_images(doc: Document, result: DOCXValidationResult):
    """Check images: embedded vs linked, inline vs floating."""
    from docx.oxml.ns import qn
    
    # Check all relationships for images
    if hasattr(doc.part, 'rels'):
        for rel in doc.part.rels.values():
            if 'image' in rel.target_ref.lower() or rel.reltype.endswith('/image'):
                # Check if it's embedded or linked
                if hasattr(rel, 'target_mode') and rel.target_mode == 'External':
                    result.summary['images_linked'] += 1
                    result.add_warning(
                        'images_linked',
                        f"Found linked image: {rel.target_ref}",
                        "ZASADA NR 1: All images must be EMBEDDED in DOCX, not linked. "
                        "In Word: File → Info → Edit Links to Files → ensure no links exist. "
                        "Right-click image → Save as Picture → re-insert as embedded."
                    )
                else:
                    result.summary['images_embedded'] += 1
    
    # Check for inline vs floating images
    body = doc.element.body
    for paragraph in body.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}p'):
        # Check for inline images
        inline_pics = paragraph.findall('.//{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}inline')
        if inline_pics:
            result.summary['images_inline'] += len(inline_pics)
        
        # Check for floating/anchored images
        anchored_pics = paragraph.findall('.//{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}anchor')
        if anchored_pics:
            result.summary['images_floating'] += len(anchored_pics)
            result.add_warning(
                'images_floating',
                f"Found {len(anchored_pics)} floating/anchored image(s)",
                "ZASADA NR 2: Use 'Inline with text' positioning. "
                "In Word: Right-click image → Wrap Text → In line with text"
            )
    
    # Check headers/footers for images
    for section in doc.sections:
        if section.header:
            header_paragraphs = section.header.paragraphs
            for para in header_paragraphs:
                if para._element.findall('.//{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}inline'):
                    result.add_warning(
                        'images_in_headers',
                        "Images found in headers/footers",
                        "ZASADA NR 3: Images in headers must be inline, without anchors, without absolute positioning"
                    )


def _validate_colors(doc: Document, result: DOCXValidationResult):
    """Check colors: RGB vs theme colors."""
    from docx.oxml.ns import qn
    
    # Check paragraph styles for theme colors
    if hasattr(doc, 'styles'):
        for style in doc.styles:
            if hasattr(style, 'font') and style.font and style.font.color:
                color = style.font.color
                if hasattr(color, 'theme_color') and color.theme_color:
                    result.summary['theme_colors'] += 1
                    result.add_warning(
                        'theme_colors',
                        f"Style '{style.name}' uses theme color: {color.theme_color}",
                        "ZASADA NR 4: Use explicit RGB colors, not Theme Colors. "
                        "In Word: Styles → Modify → Font Color → More Colors → Custom → enter RGB values"
                    )
                elif hasattr(color, 'rgb') and color.rgb:
                    result.summary['rgb_colors'] += 1
    
    # Check runs for theme colors
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            if run.font and run.font.color:
                color = run.font.color
                if hasattr(color, 'theme_color') and color.theme_color:
                    result.summary['theme_colors'] += 1
                    result.add_warning(
                        'theme_colors',
                        f"Text uses theme color: {color.theme_color}",
                        "ZASADA NR 4: Use explicit RGB colors, not Theme Colors"
                    )


def _validate_text_boxes(doc: Document, result: DOCXValidationResult):
    """Check for TextBoxes (LibreOffice doesn't render them well)."""
    body = doc.element.body
    
    # Check for text boxes (w:txbxContent)
    text_boxes = body.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}txbxContent')
    if text_boxes:
        result.summary['text_boxes'] = len(text_boxes)
        result.add_warning(
            'text_boxes',
            f"Found {len(text_boxes)} TextBox(es)",
            "LibreOffice does not render TextBoxes correctly. "
            "Remove TextBoxes and use regular paragraphs with borders/backgrounds instead."
        )


def _validate_smart_art(doc: Document, result: DOCXValidationResult):
    """Check for SmartArt (LibreOffice doesn't support it)."""
    body = doc.element.body
    
    # Check for SmartArt (usually in drawings or embedded objects)
    # SmartArt is complex to detect, but we can check for common patterns
    drawings = body.findall('.//{http://schemas.openxmlformats.org/drawingml/2006/main}graphicData')
    for drawing in drawings:
        uri = drawing.get('{http://www.w3.org/2001/XMLSchema-instance}type')
        if uri and 'smartArt' in uri.lower():
            result.summary['smart_art'] += 1
            result.add_warning(
                'smart_art',
                "Found SmartArt object",
                "LibreOffice does not support SmartArt. "
                "Convert SmartArt to regular images or shapes before conversion."
            )


def _validate_toc(doc: Document, result: DOCXValidationResult):
    """Check for TOC fields (must be updated before conversion)."""
    body = doc.element.body
    
    # Check for TOC fields (w:fldSimple or w:fldChar with TOC)
    toc_fields = []
    for para in body.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}p'):
        # Check for field simple (TOC)
        fld_simple = para.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fldSimple')
        if fld_simple is not None:
            instr = fld_simple.get(qn('w:instr'))
            if instr and 'TOC' in instr.upper():
                toc_fields.append(para)
        
        # Check for field characters (complex fields)
        fld_chars = para.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fldChar')
        for fld_char in fld_chars:
            fld_char_type = fld_char.get(qn('w:fldCharType'))
            if fld_char_type == 'begin':
                # Check if it's a TOC field by looking at following runs
                runs = para.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}r')
                for run in runs:
                    instr_text = run.text or ""
                    if 'TOC' in instr_text.upper():
                        toc_fields.append(para)
                        break
    
    if toc_fields:
        result.summary['toc_fields'] = len(toc_fields)
        result.add_warning(
            'toc_not_updated',
            f"Found {len(toc_fields)} Table of Contents field(s)",
            "ZASADA NR 7: TOC must be UPDATED before PDF conversion. "
            "In Word: Right-click TOC → Update Field → Update entire table. "
            "Headless LibreOffice does NOT update fields automatically."
        )


def _validate_cover_page(doc: Document, result: DOCXValidationResult):
    """Check for Word Cover Page (LibreOffice doesn't handle it well)."""
    # Check first section for cover page indicators
    if len(doc.sections) > 0:
        first_section = doc.sections[0]
        
        # Check if first page has different header/footer (cover page indicator)
        if hasattr(first_section, 'header') and first_section.header:
            # Check for cover page style or special formatting
            first_para = doc.paragraphs[0] if doc.paragraphs else None
            if first_para:
                style_name = first_para.style.name if first_para.style else ""
                if 'cover' in style_name.lower() or 'title' in style_name.lower():
                    # Check if it's a Word Cover Page (vs manual title page)
                    result.add_warning(
                        'cover_page',
                        "Possible Word Cover Page detected",
                        "ZASADA NR 6: Avoid Word's 'Cover Page' feature. "
                        "Use a regular first page with manual formatting instead."
                    )


def _validate_fonts(doc: Document, result: DOCXValidationResult):
    """Check fonts (warn if custom fonts might not be available)."""
    fonts_used = set()
    
    # Collect all fonts used
    if hasattr(doc, 'styles'):
        for style in doc.styles:
            if hasattr(style, 'font') and style.font and style.font.name:
                fonts_used.add(style.font.name)
    
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            if run.font and run.font.name:
                fonts_used.add(run.font.name)
    
    # Common system fonts that should be available
    common_fonts = {
        'Arial', 'Helvetica', 'Times New Roman', 'Times', 'Courier New', 'Courier',
        'Calibri', 'Verdana', 'Georgia', 'Palatino', 'Garamond'
    }
    
    # Warn about non-standard fonts
    custom_fonts = fonts_used - common_fonts
    if custom_fonts:
        result.add_warning(
            'custom_fonts',
            f"Custom fonts detected: {', '.join(custom_fonts)}",
            "ZASADA NR 5: Ensure fonts are installed on the server. "
            "LibreOffice will substitute fonts if they're not available, which may change colors and layout. "
            f"Install fonts: {', '.join(custom_fonts)}"
        )


def get_validation_summary(result: DOCXValidationResult) -> Dict:
    """Convert validation result to API response format."""
    return {
        'is_valid': result.is_valid,
        'warnings': result.warnings,
        'errors': result.errors,
        'summary': result.summary,
        'recommendations': _get_recommendations(result)
    }


def _get_recommendations(result: DOCXValidationResult) -> List[str]:
    """Generate actionable recommendations based on validation results."""
    recommendations = []
    
    if result.summary['images_linked'] > 0:
        recommendations.append("Embed all linked images before conversion")
    
    if result.summary['images_floating'] > 0:
        recommendations.append("Change floating images to 'Inline with text'")
    
    if result.summary['theme_colors'] > 0:
        recommendations.append("Convert theme colors to explicit RGB values")
    
    if result.summary['text_boxes'] > 0:
        recommendations.append("Remove TextBoxes and use regular paragraphs")
    
    if result.summary['smart_art'] > 0:
        recommendations.append("Convert SmartArt to images or shapes")
    
    if result.summary['toc_fields'] > 0:
        recommendations.append("Update Table of Contents fields in Word before conversion")
    
    if result.summary['cover_pages'] > 0:
        recommendations.append("Replace Word Cover Page with regular first page")
    
    if not recommendations:
        recommendations.append("Document appears ready for PDF conversion")
    
    return recommendations


